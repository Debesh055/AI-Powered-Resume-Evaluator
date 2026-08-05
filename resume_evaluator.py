import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel,Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")
if not  my_api_key:
    raise ValueError("api not found....!!!!")
client=Groq(api_key=my_api_key)
model="llama-3.3-70b-versatile"
role="user"
# ............................step1.......................................
job_description="""Data Analyst Job Description

 role: Data Analyst

Location: Bangalore, India (Hybrid)

Experience: 0–3 Years

Employment Type: Full-Time

Job Summary

We are looking for a detail-oriented and analytical Data Analyst to join our growing team. The ideal candidate will collect, clean, analyze, and interpret data to help business stakeholders make informed decisions. You will work closely with business, product, and engineering teams to identify trends, build dashboards, and provide actionable insights.

Key Responsibilities

Collect and analyze structured and unstructured datasets.
Automate repetitive reporting tasks using Python.
Ensure data accuracy and maintain data quality standards.
Required Skills
Strong knowledge of SQL
Python (Pandas, NumPy)
Microsoft Excel (Advanced)
Power BI or Tableau
Data Visualization
Statistics
Data Cleaning
Problem Solving
Communication Skills
Critical Thinking

Preferred Skills

Machine Learning basics
ETL concepts
Google Analytics
AWS or Azure
Git and GitHub

Qualifications

Bachelor's degree in Computer Science, Information Technology, Statistics, Mathematics, or a related field.
MCA, BCA, B.Tech, or M.Sc. candidates are welcome.

Experience

0–3 years of experience in Data Analytics or Business Intelligence.
Freshers with strong SQL and Python skills may also apply."""


class jobD(BaseModel):
    role:str
    required_skills:list[str]
    prefeared_skills:list[str]
    minimum_exprience:float | None
    educational_requirements:list[str]
    responsiblity:list[str]

job_schema=jobD.model_json_schema()

system_prompt=f"""

you are the expatr hr assistent.
your job is to analysing job description.
and give the structure  information from them.
return only valid json matching withi this schema.

{job_schema}

importent:
do not return the schema itself.
do not return the fields like "proparty","title" or "type"
fill the schema with actual information extract from the job desceription.

if miminum exprience not memction return null.
if information for a list is missing ,return an empty list.
do not give extra information

"""
user_prompt=f"""
 analyzing the following job description
 {job_description} 
"""
system_message={
    "role":"system",
    "content":system_prompt
}
user_messages={
    "role":role,
    "content":user_prompt
}
response_formate={
    "type":"json_object"
}

messages=[system_message,user_messages]
response=client.chat.completions.create(model=model,messages=messages,response_format=response_formate)
answer=response.choices[0].message.content
raw_json=answer
# print(raw_json)

import json
job_data=json.loads(raw_json)
job=jobD(**job_data)
print(job.minimum_exprience)
print(job.educational_requirements)


# resume.....................step2.....................................
class MatchResult(BaseModel):
    score:float
    details:dict
class Exprience(BaseModel):
    company:str | None = None
    role:str | None = None
    duration: str | None =None
    decription:str | None=None
    skills_used: list[str]=[]

class Resume(BaseModel):
    name:str| None=None
    email:str | None=None
    phone: str | None=None
    total_exprience: float | None=None
    skills:list[str]=[]
    exprience:list[Exprience]=[]
    education:list[str]=[]
    projects:list[str]=[]
    certifications:list[str]=[]

resume_schema=Resume.model_json_schema()
# match final result schema
def final_score(job,resume):
    match_schema=MatchResult.model_json_schema()
    prompt=f"""
      you are the hr recruiter.
      compare the candidate resume with the job decription.
      JOB DESCRIPTION:
      {job.model_dump_json(indent=2)}
      CANDIDATE RESUME:
      {resume.model_dump_json(indent=2)}

     return json matching this schema:
     {match_schema}
   give me :
   1. candidate name
   2.matching skill
   3.misssing importent skill
   4.whater a exprience is requirment is met
   5.overall matching parcentage from 0 to 100
   6.keep the final vardict why you shou choice this candidate

   keep the response consise and easy to read.
   """
    message={
        "role":"user",
        "content":prompt
    }
    messages=[message]
    response_formate={
        "type":"json_object"
    }
    response=client.chat.completions.create(model=model,messages=messages,response_format=response_formate)
    data=json.loads(response.choices[0].message.content)
    return MatchResult(**data)

#valuate the resume which is uplode in resumes folder with the help of resume class
def parse_resume(resume_text):
    system_prompt = f"""
You are an expert HR Resume Parser with extensive experience in Applicant Tracking Systems (ATS).

Your job is to convert an unstructured resume into structured JSON.

Return the extracted information using the following JSON schema:

{resume_schema}

Rules:
- Return ONLY valid JSON.
- Do not return markdown.
- Do not wrap the JSON inside code blocks.
- Do not explain your answer.
- Do not include the schema in your response.
- Do not invent any information.
- Extract only what is explicitly mentioned in the resume.
- If information is unavailable:
  • use null for scalar values.
  • use [] for list values.
- Keep names of companies, universities, certifications, and projects exactly as written.
- Include all technical and soft skills found in the resume.
- For each work experience, identify:
  • company
  • role
  • duration
  • description
  • skills_used
- If total experience cannot be calculated confidently, return null.
- Ensure the final output is valid JSON matching the schema exactly.
"""
    user_prompt=f"""
parsed the following resume:
{resume_text}
"""
    system_message={
        "role":"system",
        "content":system_prompt
    }
    user_message={
        "role":"user",
        "content":user_prompt
    }
    messages=[system_message,user_message]
    response_formate={
        "type":"json_object"
    }
    response=client.chat.completions.create(model=model,messages=messages,response_format=response_formate)
    raw_output=response.choices[0].message.content
    data=json.loads(raw_output)
    resume=Resume(**data)
    return resume

# parsced resume in text formate.............part3..............................
from pypdf import PdfReader
from docx  import Document
def read_pdf(file_path):
    reader=PdfReader(file_path)
    text=""
    for page in reader.pages:
        page_text=page.extract_text()
        if page_text:
            text +=page_text +"\n"
    return text

def read_docx(file_path):
    document=Document(file_path)
    text=""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text +"\n"
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text +=cell.text +"\n" 
    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        None

# actual resume parsed...............patr4......................................
resume_folder=Path("resumes")
all_result=[]
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in[".pdf",".docx"]:
        continue
    print("\nprocessing",file_path.name)
    resume_text=read_resume(file_path)
    parsed_resume=parse_resume(resume_text)
    time.sleep(5)
    result=final_score(job,parsed_resume)
    time.sleep(5)
    print("score",result.score)
    all_result.append({
        "name":parsed_resume.name,
        "score":result.score,
        "details":result.details
    })
all_result.sort(
    key=lambda candidate:candidate["score"],
    reverse=True
)
top_2=all_result[:2]
worst_2=all_result[-2:]
print("top 2 candidates :-")
for candidate in top_2:
    print(
        candidate["name"],"-",
        candidate["score"],"%"

    )
    print(candidate["details"])
print("worsed 2 candidates :-")
for candidate in worst_2:
    print(
        candidate["name"],"-",
        candidate["score"],"%"

    )
    print(candidate["details"])



